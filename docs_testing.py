from docx import Document

# Открываем существующий документ
doc = Document('documents/request.docx')

# Словарь замен, где ключ - текст для замены, значение - новое значение
replacements = {
    "[ka_f/l name]": "დიმიტრი ჯიოშვილი",
    "[en f/l_name]": "dimitri jioshvili",
    "[ID]": "57001060603",
    "[phone]": "571181762",
    "[tracking]": "MP0000000"
}

# Проходим по каждому параграфу в документе
for paragraph in doc.paragraphs:
    for key, value in replacements.items():
        if key in paragraph.text:
            # Заменяем текст
            paragraph.text = paragraph.text.replace(key, value)

# Сохраняем документ с новыми изменениями
doc.save('get_request.docx')



# Открываем существующий документ
doc = Document('documents/transportation invoice.docx')

# Словарь замен, где ключ - текст для замены, значение - новое значение
replacements = {
    "[date]": "14.03.2000",
    "[payer]": "dimitri jioshvili",
    "[code]": "57001060603",
    "[tracking]": "MP04873641",
    "[quantity]": "10.25",
    "[price]": "6",
}

# Функция для замены текста в ячейке
def replace_text_in_cell(cell, replacements):
    for key, value in replacements.items():
        if key in cell.text:
            cell.text = cell.text.replace(key, value)

# Проходим по каждому параграфу в документе
for paragraph in doc.paragraphs:
    for key, value in replacements.items():
        if key in paragraph.text:
            # Заменяем текст
            paragraph.text = paragraph.text.replace(key, value)

# Проходим по каждой таблице в документе
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            replace_text_in_cell(cell, replacements)
# Сохраняем документ с новыми изменениями
doc.save('get_transportation invoice.docx')