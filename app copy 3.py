from flask import Flask, render_template, request, redirect, url_for, flash, jsonify
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime, timedelta
from flask_login import login_required, login_user, logout_user, current_user
import os 
from models import Purcell, db, User, login_manager, Menu, Storage, Booking, Forms, Expertise
from config import secret_key
from sqlalchemy import func, desc, update
from flask.views import View, MethodView
from flask import send_file
from openpyxl.styles import Alignment, Border, Font, Side
from openpyxl import Workbook, load_workbook
from sqlalchemy import or_
import asyncio
import openpyxl
import threading
from bot import send_location_message
import asyncio

from PIL import Image
from PIL.ExifTags import TAGS
from sqlalchemy.exc import SQLAlchemyError
import re
from openpyxl.drawing.image import Image
from decimal import Decimal
from docx import Document

import json

import random
from io import BytesIO

import qrcode


from functions import  get_last_record, calculate_cost, handle_image, handle_uploaded_image, get_reservation_data, validate_input, format_trecing, save_record, generate_new_number, add_record, get_sorted_dates, update_json_file, delete_old_data, clean_old_files, log_error, process_payment, save_booking_to_db, get_existing_booking, update_booking, manifest_filter, xml_convertor, allowed_file, trecing_redactor, edit_parcel_


app = Flask(__name__)



app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+mysqlconnector://root:QazEdcQweZxcQscEsz123@localhost/packages'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False  # Отключает отслеживание изменений
app.config['REMEMBER_COOKIE_DURATION'] = timedelta(hours=8)
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=8)
app.config['UPLOAD_FOLDER'] = os.path.join(os.getcwd(), 'static/purcells')
app.config['SECRET_KEY'] = secret_key

app.config.from_object(__name__)

db.init_app(app)

login_manager.init_app(app)
login_manager.login_view = 'login'



#-------------------------------------------------------------------------------------------------#
# ------------------------------               /login               ------------------------------#
   
class LoginViev(MethodView):

    def get(self):
        return render_template('login.html')
    


    def post(self):

        # Получение данных из формы
        login_input = request.form.get('login')
        password = request.form.get('psw')

        # Поиск пользователя по введенному логину
        user = User.query.filter_by(login=login_input).first()

        # Проверка существования пользователя и совпадения пароля
        if user and user.check_password(password):
            # Если логин и пароль совпадают, авторизуем пользователя
            login_user(user)
            return redirect(url_for('index'))
        else:
            # Если логин или пароль неверные, выводим сообщение об ошибке
            flash('მომხმარებელი არ მოიძებნა ან პაროლი არ ემთხვევა', category='error')
        return render_template('login.html')


app.add_url_rule('/login', view_func=LoginViev.as_view('login'))
#-------------------------------------------------------------------------------------------------#
# ------------------------------               /index               ------------------------------#
   
class IndexViev(MethodView):
    decorators = [login_required]

    def get(self):
        try:
            # Получаем и сортируем даты
            msk_dates, spb_dates = get_sorted_dates(db)
            # Обновляем JSON-файл с датами
            update_json_file(msk_dates, spb_dates)
            # Удаляем устаревшие данные
            delete_old_data(db)
            # Очищаем старые файлы
            clean_old_files(app.config['UPLOAD_FOLDER'])
        except Exception as e:
            # Если произошла ошибка, логируем её
            log_error(e)
        
        # Возвращаем HTML-шаблон с данными для отображения на странице
        return render_template('index.html', msk_dates=msk_dates, spb_dates=spb_dates)  


# Регистрация классов в приложении
app.add_url_rule('/', view_func=IndexViev.as_view('index'))

#-------------------------------------------------------------------------------------------------#
# ------------------------------               /all                 ------------------------------#

class AllView(MethodView):
    decorators = [login_required]

    def get(self):
        # Получаем текущую дату
        today = datetime.now().date()
        delta = timedelta(days=60)
        date_threshold = today - delta

        # Получаем данные для отображения, отсортированные по дате и номеру
        all_data = list(reversed(Purcell.query.filter(Purcell.date >= date_threshold)
                         .order_by(Purcell.date.asc(), Purcell.number.asc())
                         .all()))

        # Получаем даты последних 10 рейсов для отображения из уже полученных данных
        last_10_flights = list(set(row.flight for row in all_data))[:10]

        # Отображаем шаблон страницы 'all.html' с данными
        return render_template('all.html', all_data=all_data, last_10_flights=last_10_flights)


class RemoveFromListView(MethodView):
    decorators = [login_required]

    def post(self):
        access = ['admin', 'Tbilisi']
        # Проверка прав доступа
        if current_user.role not in access:
            return jsonify({'success': False, 'message': 'თქვენ არ გაქვთ წვდომა'}), 404
        
        # Получаем значение data.id из запроса
        data_id = request.json.get('id')

        # Ищем запись в таблице Purcell по переданному id
        purcell_entry = Purcell.query.get(data_id)

        if purcell_entry:
            # Если запись найдена, удаляем ее
            db.session.delete(purcell_entry)
            db.session.commit()
            return jsonify({'success': True, 'message': 'ჩანაწერი წაშლილია'}), 200
        else:
            # Если запись не найдена, возвращаем сообщение об ошибке
            return jsonify({'success': False, 'message': f'Запись с id {data_id} не найдена'}), 404
        

class EditParcelView(MethodView):
    decorators = [login_required]

    def post(self):
        try:
            # Получаем данные из формы
            data = request.form.to_dict()

            # Получаем переданную фотографию, если она есть
            photo = request.files.get('photo')

            # Обработка фотографии, если она была передана
            if photo and photo.filename != '':
                handle_uploaded_image(request.files['photo'], data['id'], app)

            edit_parcel_(db, data)

            # Возвращаем сообщение об успешной обработке
            return jsonify({'message': 'რედაქტირებამ წარმატებით ჩაიარა', 'success': True}), 200
        except Exception as e:
            return jsonify({'message': f'დაფიქსირდა შეცდომა : {e}', 'success': False}), 400




class DocumentsView(MethodView):
    decorators = [login_required]

    def get(self):
        access = ['admin', 'Tbilisi']
        if current_user.role not in access:
            flash('თქვენ არ გაქვთ წვდომა ამ გვერდზე', 'error')
            return redirect(url_for('index'))
    
        return render_template('documents.html')


class CreateDocxRequestView(MethodView):
    decorators = [login_required]

    def post(self):
        # Получаем данные из формы
        ka_name = request.form['ka_f-l_name']
        ge_name = request.form['ge_f-l_name']
        personal_id = request.form['id']
        phone = request.form['phone']
        tracking = request.form['tracking']

        # Открываем существующий документ
        doc = Document('documents/request.docx')

        # Словарь замен
        replacements = {
            "[ka_f/l name]": ka_name,
            "[en f/l_name]": ge_name,
            "[ID]": personal_id,
            "[phone]": phone,
            "[tracking]": f'MP{tracking}'
        }

        # Проходим по каждому параграфу в документе
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    # Заменяем текст
                    paragraph.text = paragraph.text.replace(key, value)

        # Сохраняем документ в память (BytesIO)
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Отправляем документ на скачивание
        file_name = f'{ka_name}.docx'
        return send_file(file_stream, as_attachment=True, download_name=file_name)



class CreateDocxGetTransportingInvoiceView(MethodView):
    decorators = [login_required]

    def post(self):
        # Получаем данные из формы
        date = request.form['date']
        payer = request.form['payer']
        code = request.form['code']
        tracking = request.form['tracking']
        quantity = request.form['quantity']
        price = request.form['price']
        broker_service = request.form.get('broker_service', 'No')

        tot_transp = int(price) * int(quantity)

        if broker_service == 'on':
            tot_pri = tot_transp + 15

            replacements = {
            "[date]": date,
            "[payer]": payer,
            "[code]": code,
            "[tracking]": f'MP{tracking}',
            "[quantity]": quantity,
            "[price]": price,
            "[tot_transp]": f"{tot_transp:.2f}",
            "[tot_pri]": f"{tot_pri:.2f}",
            "[2]": "2",
            "[საბროკერო მომსახურეობა]": "საბროკერო მომსახურეობა",
            "[ერთ]": "ერთ",
            "[1]": "1",
            "[15]": "15"
        }
        else:
            tot_pri = tot_transp

            replacements = {
            "[date]": date,
            "[payer]": payer,
            "[code]": code,
            "[tracking]": f'MP{tracking}',
            "[quantity]": quantity,
            "[price]": price,
            "[tot_transp]": f"{tot_transp:.2f}",
            "[tot_pri]": f"{tot_pri:.2f}",
            "[2]": "",
            "[საბროკერო მომსახურეობა]": "",
            "[ერთ]": "",
            "[1]": "",
            "[15]": ""
        }
        # Открываем существующий документ
        doc = Document('documents/transportation invoice.docx')


        # Функция для замены текста в ячейке
        def replace_text_in_cell(cell, replacements):
            for key, value in replacements.items():
                if key in cell.text:
                    cell.text = cell.text.replace(key, value)

        # Проходим по каждому параграфу в документе
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    # Заменяем текст
                    paragraph.text = paragraph.text.replace(key, value)

        # Проходим по каждой таблице в документе
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_text_in_cell(cell, replacements)

        # Сохраняем документ в память (BytesIO)
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Отправляем документ на скачивание
        file_name = f'Transporting_Invoice_{payer}.docx'
        return send_file(file_stream, as_attachment=True, download_name=file_name)


app.add_url_rule('/documents', view_func=DocumentsView.as_view('documents'))
app.add_url_rule('/create_docxrequest', view_func=CreateDocxRequestView.as_view('create_docxrequest'))
app.add_url_rule('/create_transporting_invoice', view_func=CreateDocxGetTransportingInvoiceView.as_view('create_transporting_invoice'))



class EditDeliveryViev(MethodView):
    decorators = [login_required]

    def post(self):
        access = ['admin', 'Moscow', 'SPB']
        # Проверка прав доступа
        if current_user.role not in access:
            return jsonify({'message': 'თქვენ არ გაქვთ წვდომა!', 'success': False}), 404
        
        data_id = request.json.get('id')  # Получаем ID из запроса

        # Находим запись в таблице Purcell по переданному ID
        purcell_entry = Purcell.query.get(data_id)

        if purcell_entry:
            # Проверяем, что статус доставки еще не 'yes'
            if purcell_entry.delivery != 'yes':
                # Изменяем статус доставки на 'yes'
                purcell_entry.delivery = 'yes'
                db.session.commit()
                return jsonify({'message': 'Посылка вручена', 'success': True}), 200
            else:
                # Если статус доставки уже 'yes', возвращаем ошибку 404
                return jsonify({'message': 'Данная посылка уже вручена!', 'success': False}), 404
        else:
            return jsonify({'message': 'Запись не найдена'}), 404


# Регистрация классов в приложении
app.add_url_rule('/all', view_func=AllView.as_view('all'))
app.add_url_rule('/removing_from_the_list', view_func=RemoveFromListView.as_view('removing_from_the_list'))
app.add_url_rule('/edit_parcel', view_func=EditParcelView.as_view('edit_parcel'))
app.add_url_rule('/delivery_status', view_func=EditDeliveryViev.as_view('delivery_status'))

#-------------------------------------------------------------------------------------------------#

# ------------------------------               /all end/            ------------------------------#


#-------------------------------------------------------------------------------------------------#
# ------------------------------               /add                 ------------------------------#

class AddParcel(MethodView):
    decorators = [login_required]

    def get(self):
        access = ['admin', 'Tbilisi', 'Batumi']
        if current_user.role not in access:
            flash('თქვენ არ გაქვთ წვდომა ამ გვერდზე', 'error')
            return redirect(url_for('index'))
    
        return render_template('add.html')



class SaveParcel(MethodView):
    decorators = [login_required]

    def post(self):
        data = request.form.to_dict()
        last_record = get_last_record(db)

        new_record = generate_new_number(data, last_record, db)
        cost = calculate_cost(request.form.get('payment'), request.form.get('cost'), request.form.get('payment_currency'))
        
        try:
            handle_image(request.files['photo'], new_record, data['currentDateTime'], app)
        except Exception as e:
            # Обработка ошибки и возврат сообщения об ошибке
            response_data = {'success': False, 'message': f'დაფიქსირდა შეცდომა ფოტოს შენახვისას: {str(e)}'}
            return jsonify(response_data)
        
        try:
            add_record(new_record, data, cost, db, current_user.role)
        except Exception as e:
            # Обработка ошибки при сохранении записи и возврат сообщения об ошибке
            response_data = {'success': False, 'message': f'Ошибка при сохранении записи: {str(e)}'}
            return jsonify(response_data)
        
        response_data = {
            'success': True,
            'message': f'ამანათი წარმატებიით დაემატა! № {new_record}'
        }
        return jsonify(response_data)




# Регистрация классов в приложении
app.add_url_rule('/add', view_func=AddParcel.as_view('add'))
app.add_url_rule('/saving_a_parcel', view_func=SaveParcel.as_view('saving_a_parcel'))

#-------------------------------------------------------------------------------------------------#
# ------------------------------               /storage             ------------------------------#


class StorageViev(MethodView):
    decorators = [login_required]

    def get(self):
        access = ['admin', 'Tbilisi']
        if current_user.role not in access:
            flash('თქვენ არ გაქვთ წვდომა ამ გვერდზე', 'error')

            return redirect(url_for('index'))
    
        return render_template('storage.html')


class StorageSaveViev(MethodView):
    decorators = [login_required]

    def post(self):
        try:
            shelf = request.form.get('shelf')
            trecing = request.form.get('trecing')

            trecing_list = trecing.split()

            if not validate_input(shelf, trecing):
                return jsonify({'success': False, 'message': 'დამატება ვერ მოხერხდა: შეავსეთ მოცემული ველები!'})
            
            date = datetime.now().date()
            last_shelf = None  # Переменная для отслеживания последнего значения
            for record in trecing_list:
                formatted_trecing = format_trecing(record)

                try:
                    last_shelf = save_record(shelf, formatted_trecing, date, db)
                except SQLAlchemyError as e:
                    return jsonify({'success': False, 'message': 'დაიკარგა მონაცემთა ბაზასთან კავშირი!'})

            return jsonify({'last': last_shelf})
        except Exception as e:
            return jsonify({'error': 'მოხდა ამოუცნობი შეცდომა.'})


class StorageFindView(MethodView):
    decorators = [login_required]

    def post(self):
        trecing = request.form['trecing']
        info = request.form['info']
        # Выполняем поиск посылки в базе данных
        storage = Storage.query.filter_by(trecing=trecing).first()
        if storage:
            location = storage.shelf  # Местоположение посылки
            asyncio.run_coroutine_threadsafe(send_location_message(trecing, location, info, storage.date), loop)
            if info:
                db.session.delete(storage)
                db.session.commit()
        else:
            location = "ამანათი არ მოიძებნა"
        
        # Возвращаем данные в формате JSON
        return jsonify({'shelf': location})


app.add_url_rule('/find', view_func=StorageFindView.as_view('find'))

def start_loop(loop):
    asyncio.set_event_loop(loop)
    loop.run_forever()

loop = asyncio.new_event_loop()
t = threading.Thread(target=start_loop, args=(loop,))
t.start()



@app.route('/user_add', methods=['POST'])
@login_required
def add_user():
    user_id = request.json.get('user_id')

    # Открываем файл bot_users.json и загружаем его содержимое
    try:
        with open('bot_users.json', 'r') as file:
            data = json.load(file)
    except FileNotFoundError:
        data = []

    # Проверяем, что user_id является уникальным числовым значением
    if user_id.isdigit() and int(user_id) not in data:
        data.append(int(user_id))
        # Сохраняем обновленные данные в файл bot_users.json
        with open('bot_users.json', 'w') as file:
            json.dump(data, file)
        response_data = {"message": "Пользователь добавлен успешно"}
    else:
        response_data = {"message": "Недопустимый или дублирующийся пользователь"}

    return jsonify(response_data)


app.add_url_rule('/storage', view_func=StorageViev.as_view('storage'))
app.add_url_rule('/save', view_func=StorageSaveViev.as_view('save'))
#-------------------------------------------------------------------------------------------------#
# ------------------------------               /images_list         ------------------------------#


@app.route('/images_list', methods=['POST'])
def parcell_pictures_list():
    start_date = request.form['startDate']
    end_date = datetime.strptime(request.form['endDate'], '%Y-%m-%d')
    end_date += timedelta(days=1)
    city = request.form.get('images_city')


    # Запрос к базе данных для выбора записей в заданном диапазоне дат
    purcells = Purcell.query.filter(Purcell.date.between(start_date, end_date), Purcell.delivery == 'no', Purcell.city == city).all()

    
    return render_template('images_list.html', purcells=purcells, start_date=start_date, end_date=end_date)


@app.route('/images_list_delivery_status', methods=['POST'])
def images_list_delivery_status():
    data_id = request.json.get('id')
    purcell_entry = Purcell.query.get(data_id)

    if purcell_entry:
        if purcell_entry.delivery == 'yes':
            return jsonify({'message': 'ამანათი უკვე გაცემული იყო!', 'success': False}), 400
        
        purcell_entry.delivery = 'yes'
        db.session.commit()
        return jsonify({'message': 'ამანათი გაცემულია', 'success': True}), 200
    else:
        return jsonify({'message': 'ამანათი არ მოიძებნა', 'success': False}), 400
    




#-------------------------------------------------------------------------------------------------#
# ------------------------------               /reservation         ------------------------------#



@app.route('/reservation', methods=['POST', 'GET'])
@login_required
def reservation():

    access = ['admin', 'Tbilisi', 'Moscow']
    if current_user.role not in access:
        flash('თქვენ არ გაქვთ წვდომა ამ გვერდზე', 'error')
        return redirect(url_for('index'))
    
    selected_date = request.args.get('date')
    reis = request.args.get('route')
    reservation_data = get_reservation_data(selected_date, reis, 55)
    return render_template('reservation.html', **reservation_data)



@app.route('/reservation_big', methods=['POST', 'GET'])
@login_required
def reservation_big():
    access = ['admin', 'Tbilisi', 'Moscow']
    if current_user.role not in access:
        flash('თქვენ არ გაქვთ წვდომა ამ გვერდზე', 'error')
        return redirect(url_for('index'))
    selected_date = request.args.get('date')
    reis = request.args.get('route')
    reservation_data = get_reservation_data(selected_date, reis, 59)
    return render_template('reservation_big.html', **reservation_data)



@app.route('/save_data', methods=['POST'])
@login_required
def save_data():
    if request.method == 'POST':
        print(request.form.get('date_of_birth'))
        # Получаем данные из запроса
        selected_date = request.form.get('selected_date')
        seat_number = request.form.get('seat_number')
        flname = request.form.get('flname').upper()
        gender = request.form.get('gender')
        phone = request.form.get('phone')
        pasport = request.form.get('pasport')
        comment = request.form.get('comment')
        payment = request.form.get('payment')
        fwc = request.form.get('reis')
        destination = request.form.get('destination')
        pay = request.form.get('payment_method')
        pay_method = request.form.get('payment_method_card')
        date_of_birth = request.form.get('date_of_birth')
        # Обработка оплаты
        payment = process_payment(payment, pay, pay_method)
        
        # Проверяем, существует ли уже запись с таким flname
        existing_booking = Booking.query.filter_by(flname=flname, data=selected_date, fwc=fwc).first()
        if existing_booking:
            return jsonify({'success': False, 'message': 'ამ სახელსა და გვარზე ადგილი უკვე დაჯავშნილია!!!!'}), 400
        else:
            # Сохранение данных в базе данных
            save_booking_to_db(db, selected_date, seat_number, flname, gender, phone, pasport, comment, payment, fwc, destination, date_of_birth)

            return jsonify({"message": "Данные успешно сохранены"})



@app.route('/edit_booking', methods=['POST'])
@login_required
def edit_booking():
    if request.method == 'POST':
        seat_number = request.form.get('seat_number')
        old_seat_number = request.form.get('old_seat_number')  # Получаем старое место
        gender = request.form.get('gender')
        flname = request.form.get('flname').upper()
        phone = request.form.get('phone')
        pasport = request.form.get('pasport')
        comment = request.form.get('comment')
        payment = request.form.get('payment')
        destination = request.form.get('destination')
        reis = request.form.get('reis')
        selected_date = request.form.get('selected_date')
        pay = request.form.get('payment_method')
        pay_method = request.form.get('payment_method_card')
        date_of_birth = request.form.get('date_of_birth')
        payment = process_payment(payment, pay, pay_method)
        
        existing_booking = get_existing_booking(reis, selected_date, old_seat_number)
        existing_booking_old = get_existing_booking(reis, selected_date, seat_number)
        
        if seat_number == old_seat_number:
            if existing_booking_old:
                update_booking(db, existing_booking_old, gender, flname, phone, pasport, payment, destination, comment, seat_number, date_of_birth)
                return jsonify({'success': True, 'message': 'message'})

        if existing_booking:
            return jsonify({'success': False, 'message': 'ადგილი დაკავებულია'})
        elif existing_booking is None:
 
            update_booking(db, existing_booking_old, gender, flname, phone, pasport, payment, destination, comment, old_seat_number, date_of_birth)
            return jsonify({'success': True, 'message': 'message'})








@app.route('/booking_del', methods=['POST'])
@login_required
def booking_del():
    if request.method == 'POST':
        seat_number = request.form.get('s_n')
        reis = request.form.get('reis')
        selected_date = request.form.get('selected_date')
        
        existing_booking = Booking.query.filter_by(fwc=reis, data=selected_date, position=seat_number).first() #новое место
        if existing_booking:

            db.session.delete(existing_booking)
            db.session.commit()

            
        return jsonify({'success': True, 'message': 'Бронирование успешно удалено'})





@app.route('/download_ved', methods=['POST'])
@login_required
def download_ved():
    if current_user.role == 'Moscow':
        destination_address = 'Тбилиси'
    elif current_user.role == 'Tbilisi':
        destination_address = 'Москва'
    else:
        destination_address = ''
    reis = request.form.get('reis')
    selected_date = request.form.get('selected_date')
    filtered_data = Booking.query.filter(
        Booking.fwc == reis,
        Booking.data == selected_date,
        Booking.action == 'yes').all()

    workbook = openpyxl.load_workbook('vedom.xlsx')
    sheet = workbook.active

    start_row = 7
    col_letters = ['B', 'D', 'E', 'F', 'G']

    # Устанавливаем стиль шрифта по умолчанию с размером 10
    font = Font(size=10)
    for row_idx, data in enumerate(filtered_data, start=start_row):
        for col_idx, col_letter in enumerate(col_letters):
            cell = sheet[f'{col_letter}{row_idx}']
            if col_letter == 'B':
                cell.value = data.date_of_birth
            elif col_letter == 'D':
                cell.value = data.pasport
            elif col_letter == 'E':
                cell.value = data.flname
            elif col_letter == 'F':
                cell.value = data.position
            elif col_letter == 'G':
                cell.value = destination_address
            cell.font = font  # Применяем стиль шрифта к ячейке

    new_filename = 'Ведомость.xlsx'
    workbook.save(new_filename)
    workbook.close()

    return send_file(new_filename, as_attachment=True)





def apply_styles_to_cell(sheet, cell, value):
    bold_font = Font(bold=True)
    sheet[cell] = value
    sheet[cell].alignment = Alignment(horizontal='center', vertical='center')
    sheet[cell].font = bold_font



@app.route('/ticket', methods=['POST'])
@login_required
def generate_ticket():
    try:
        s_n = request.form.get('s_n')
        reis = request.form.get('reis')
        selected_date = request.form.get('selected_date')

        # Загружаем шаблон Excel-файла
        template_path = 'ticket.xlsx'
        workbook = load_workbook(template_path)
        sheet = workbook.active

        # По вашим указаниям, заполняем ячейки данными
        booking = Booking.query.filter(
            Booking.fwc == reis,
            Booking.data == selected_date,
            Booking.position == s_n
        ).first()
       

        if booking:
            booking.action = 'yes'
            db.session.commit()  # Сохранение изменений в базе данных
            bold_font = Font(bold=True)

            apply_styles_to_cell(sheet, 'G5', booking.flname)
            
            apply_styles_to_cell(sheet, 'S5', booking.pasport)

            apply_styles_to_cell(sheet, 'N8', booking.data)

            apply_styles_to_cell(sheet, 'W8', booking.destination)

            apply_styles_to_cell(sheet, 'A12', booking.position)

            if current_user.role == 'Moscow':
                sheet['A15'] = '10 : 00'
            else:
                sheet['A15'] = '10 : 00'
            sheet['A15'].alignment = Alignment(horizontal='center', vertical='center')
            sheet['A15'].font = bold_font


            if booking.gender == 'male':
                sheet['G8'] = 'М / მმ'
            elif booking.gender == 'female':
                sheet['G8'] = 'Ж / მდ'
            sheet['G8'].alignment = Alignment(horizontal='center', vertical='center')
            sheet['G8'].font = bold_font


            # Обработка оплаты
            payment_value = booking.payment
            payment_currency = payment_value[-3:]
            if payment_currency == 'RUB' and current_user.role == 'Moscow':
                sheet['A7'] = '₽'
                sheet['A7'].alignment = Alignment(horizontal='center', vertical='center')
                sheet['A7'].font = bold_font
                sheet['D7'] = re.sub(r'\D', '', payment_value)  # Убираем первую букву
                sheet['D7'].alignment = Alignment(horizontal='center', vertical='center')
                sheet['D7'].font = bold_font
            elif payment_currency == 'GEL':
                sheet['A7'] = '₾'

                sheet['A7'].alignment = Alignment(horizontal='center', vertical='center')
                sheet['A7'].font = bold_font
                sheet['D7'] = re.sub(r'\D', '', payment_value)  # Убираем первую букву
                sheet['D7'].alignment = Alignment(horizontal='center', vertical='center')
                sheet['D7'].font = bold_font
            else:
                sheet['A7'] = '₾'
                sheet['A7'].alignment = Alignment(horizontal='center', vertical='center')
                sheet['A7'].font = bold_font
                sheet['D7'] = '200'
                sheet['D7'].alignment = Alignment(horizontal='center', vertical='center')
                sheet['D7'].font = bold_font
                
 

            timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
            data = f"კომპანია VIP-TOUR, მისამართი: გულიას ქუჩა №5. ბილეთის მიღების დრო {timestamp}. ბოლო გაჩერება-ოფისი: იუჟნაპორტოვაიას ქუჩა, სახლი 7, შენობა 25.\nКомпания VIP-TOUR, Адрес:  ул.Гулия №5.  / Время получения билета {timestamp}.  / конечная остановка-офис: ул.Южнопоровая дом 7, строение 25."
            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_L,
                box_size=6.8,
                border=0,
            )
            qr.add_data(data)
            qr.make(fit=True)

            img = qr.make_image(fill_color="black", back_color="white")
            img.save("qr_code.png")


            # Вставьте QR-код в Excel
            img = Image("qr_code.png")
            img.width = img.width // 3.2  # Уменьшите размер изображения, если необходимо
            img.height = img.height // 3.2
            sheet.add_image(img, "X10")


            # Сохраняем Excel в память
            output = BytesIO()
            workbook.save(output)
            output.seek(0)

            # Отправляем файл для скачивания
            return send_file(output, as_attachment=True, download_name='bileti.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

        else:
            return jsonify({'success': False, 'message': 'Данные не найдены'})

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})


#-------------------------------------------------------------------------------------------------#
# ------------------------------               /list                ------------------------------#


@app.route('/list', methods=['POST', 'GET'])
@login_required
def blanks():
    date_param = request.args.get('date')  # Получите значение параметра "date"
    city_param = request.args.get('where_from')  # Получите значение параметра "city"


    data = Forms.query.filter_by(date=date_param, where_from=city_param).order_by(desc(Forms.number)).all()
    list_delivery = [data for data in data if data.address]

    data_dict = {
        'GEL': {'paid': 0, 'card': 0, 'not_paid': 0},
        'RUB': {'paid': 0, 'card': 0, 'not_paid': 0},
        'USD': {'paid': 0, 'card': 0, 'not_paid': 0},
        'EUR': {'paid': 0, 'card': 0, 'not_paid': 0}
    }
    total_weight = 0
    for record in data:
        try:
            weights_data = record.weights  # Получение значения поля weights из записи
            numbers = [Decimal(num) for num in weights_data.split()]
            total_weight += sum(numbers)  # Суммирование чисел
        except Exception as e:
            total_weight = ' [ ERROR !! ] '


    for item in data:
        currency = item.currency
        payment_status = item.payment_status
        price = int(item.cost)

        data_dict[currency][payment_status] += price

    # Теперь у вас есть словарь, в котором данные разделены по валюте и статусу оплаты.


    gel_paid = data_dict['GEL']['paid']
    gel_card = data_dict['GEL']['card']
    gel_not_paid = data_dict['GEL']['not_paid']
    
    rub_paid = data_dict['RUB']['paid']
    rub_card = data_dict['RUB']['card']
    rub_not_paid = data_dict['RUB']['not_paid']

    usd_paid = data_dict['USD']['paid']
    usd_card = data_dict['USD']['card']
    usd_not_paid = data_dict['USD']['not_paid']

    eur_paid = data_dict['EUR']['paid']
    eur_card = data_dict['EUR']['card']
    eur_not_paid = data_dict['EUR']['not_paid']

    # Здесь вы можете использовать отфильтрованные данные для вашей веб-страницы
    return render_template('list.html', data=data, gel_paid=gel_paid, gel_card=gel_card, gel_not_paid=gel_not_paid,
                           rub_paid=rub_paid, rub_card=rub_card, rub_not_paid=rub_not_paid,
                           usd_paid=usd_paid, usd_card=usd_card, usd_not_paid=usd_not_paid,
                           eur_paid=eur_paid, eur_card=eur_card, eur_not_paid=eur_not_paid,
                           total_weight=total_weight, city_param=city_param, date_param=date_param,
                           list_delivery=list_delivery)



@app.route('/add_parcell_to_list', methods=['POST'])
@login_required
def add_parcell_to_list():
    access = ['admin', 'Moscow', 'SPB']
    if current_user.role not in access:
        return 'error user'
    try:
        data = request.form.to_dict()

        highest_number = db.session.query(func.max(Forms.number)).filter(Forms.date == data['date'],
                                                                     Forms.where_from == data['where_from']).scalar()
        if highest_number is not None:
            new_number = highest_number + 1
        else:
            new_number = 1
        # Получите данные из JSON-запроса
        if data['cost'] == '':
            cost = 0
        else:
            cost = data['cost']

        if data['passport'] == '':
            passport = '---'
        else:
            passport = data['passport']
        
        new_parcel = Forms(
            number = new_number,
            date=data['date'],
            sender_fio = data['sender_fl'].upper(),
            sender_phone=data['sender_phone'],
            recipient_fio=data['recipient_fl'].upper(),
            recipient_phone=data['recipient_phone'],
            passport= passport,
            city=data['city'],
            comment=data['comment'],
            price=int(cost),
            weights=data['weights'],
            cost=int(data['payment']),
            payment_status=data['payment_status'],
            currency=data['payment_currency'],
            where_from=data['where_from'],
            address = data['address']
            )

        db.session.add(new_parcel)
        db.session.commit()


            
        return redirect(url_for('blanks', date=request.form['date'], where_from=request.form['where_from']))

    except Exception as e:
        # В случае ошибки верните сообщение об ошибке
        response = {"error": str(e)}
        return response


@app.route('/list_edit_id', methods=['GET', 'POST'])
@login_required
def list_edit():
    if request.method == 'GET':
        item_id = request.args.get('id')  # Получаем id из параметра запроса
        data = Forms.query.get(item_id)
        if data is None:
            pass
        return render_template('list_edit.html', data=data)
    elif request.method == 'POST':
        item_id = request.form.get('id')
        parcel = Forms.query.get(item_id)

        parcel.sender_fio = request.form.get('sender_fl').upper()
        parcel.sender_phone = request.form.get('sender_phone')
        parcel.recipient_fio = request.form.get('recipient_fl').upper()
        parcel.recipient_phone = request.form.get('recipient_phone')
        parcel.passport = request.form.get('passport')
        parcel.city = request.form.get('city')
        parcel.comment = request.form.get('comment')
        parcel.price = int(request.form.get('cost'))
        parcel.weights = request.form.get('weights')
        parcel.cost = int(request.form.get('payment'))
        parcel.payment_status = request.form.get('payment_status')
        parcel.currency = request.form.get('payment_currency')
        parcel.address = request.form.get('address')
        db.session.commit()
        return redirect(url_for('blanks', date=request.form.get('date'), where_from=request.form.get('where_from')))


@app.route('/list_delete', methods=['POST'])
@login_required
def list_delete():
    if request.method == 'POST':
        item_id = request.form.get('item_id')
        parcel = Forms.query.get(item_id)
        if parcel:  # Проверяем, существует ли запись с таким ID
            db.session.delete(parcel)  # Удаляем запись из сессии
            db.session.commit()         # Применяем изменения в базе данных
            return redirect(url_for('blanks', date=request.form.get('date'), where_from=request.form.get('where_from')))
        # Здесь добавьте код для удаления записи из базы данных по item_id
        return redirect(url_for('blanks', date=request.form.get('date'), where_from=request.form.get('where_from')))


def random_names():
    names = [
    "ALEXANDER IVANOV", "EKATERINA PETROVA", "DMITRY SMIRNOV", "ANNA FEDOROVA", "SERGEI KOZLOV",
    "MARIA NOVIKOVA", "IVAN SOKOLOV", "ANASTASIA KUZNETSOVA", "ARTEM POPOV", "OLGA MOROZOVA",
    "VLADIMIR KUZMIN", "IRINA PAVLOVA", "NIKOLAI ZAKHAROV", "YULIA SEMENOVA", "PAVEL KONDRATOV",
    "SVETLANA SMIRNOVA", "ANDREI FEDOROV", "ELENA VASILIEVA", "ANDREY LEBEDEV", "TATIANA KUZNETSOVA",
    "ALEXEY MEDVEDEV", "NATALIA YAKOVLEVA", "VIKTORIA PROKOPIEVA", "MIKHAIL SOKOLOV", "ANGELINA KONOVALENKO",
    "RUSLAN VORONIN", "YANA PETROVA", "IGOR KARPOV", "VALERIA STEPANOVA", "ANTONINA MALININA", "KONSTANTIN ZAITSEV",
    "MARINA ROMANOVA", "PAVEL SERGEYEV", "OLGA KIRILLOVA", "ILYA KUZNETSOV", "KSENIA MOROZOVA", "DENIS POPOV",
    "VICTORIA KOVALENKO", "YURY ANTONOV", "JULIA KOROLEVA", "ALEXANDRA SIDOROVA", "MAXIM KUZMIN", "LARISA PAVLOVA",
    "SERGEY GORSHKOV", "ANNA KONONOVA", "ALEKSEI FOMIN", "EKATERINA KOROLEVA", "ARSENII VOLKOV", "IRINA ZAKHAROVA",
    "ALEKSANDR GORBUNOV", "ALEXANDRA KUZNETSOVA", "ANDREI MOROZOV", "VICTORIA SEMENOVA", "MAXIM FEDOROV",
    "OLGA PETROV", "KIRILL KONDRATOV", "YULIA ROMANOVA", "DENIS KOVALENKO", "SVETLANA ZAKHAROVA",
    "ANTONINA SIDOROVA", "DMITRY LEBEDEV", "EKATERINA SERGEYEVA", "ILYA KIRILLOV", "MARIA GORBUNOVA",
    "IGOR MALININ", "ANNA KUZMINA", "ARTEM ZAITSEV", "ELENA KONOVA", "NIKOLAI GORSHKOV", "VALERIA PAVLOVA",
    "SERGEY PROKOPIEV", "ANGELINA MEDVEDEVA", "VLADIMIR ANTONOV", "TATIANA KOROLEVA", "ANDREY VORONIN",
    "LARISA YAKOVLEVA", "PAVEL KARPOV", "NATALIA KIRKOROVA", "MIKHAIL ZAITSEV", "KSENIA SOKOLOVA",
    "YURY KUZNETSOV", "MARINA KUZMINA", "ALEKSANDR FOMIN", "ELENA VORONOVA", "ALEXEI KIRILLOV",
    "VALERIYA KUZNETSOVA", "VIKTOR KOROLEV", "ANNA SMIRNOVA", "ANDREY SOKOLOV", "OLGA GORBUNOVA",
    "ALEKSANDRA ROMANOVA", "DMITRIY KUZMIN", "EKATERINA ZAKHAROVA", "MAXIM KARPOV", "YANA KONONOVA",
    "VLADIMIR LEBEDEV", "MARIYA KIRKOROVA", "ANDREY PETUKHOV", "DARIA SMOLYAKOVA", "ALEXEY KONDRATENKO", 
    "SVETLANA IVANOVA", "IGOR SEMYONOV", "MARINA LUKINA", "ANDREI SOKOLOVSKY", "EKATERINA KAZAKOVA",
    "VLADIMIR PETROVICH", "ANNA DUBROVSKAYA", "DMITRY STEPANOV", "ELENA FEDOTOVA", "ALEXANDER ROMANOV",
    "OLGA GAVRILOVA", "MAXIM TARASOV", "YULIA KORNEEVA", "NIKOLAI SHIROKOV", "MARIA TIMOFEEVA",
    "PAVEL VASILIEV", "EKATERINA KOSHKINA", "DMITRIY MOROZOV", "NATALIA KOVALEVSKAYA", "ANDREY KOZLOV",
    "SVETLANA EGOROVA", "SERGEY BORISOV", "ANASTASIA EGOROVA", "ALEXANDER BELIAKOV", "EKATERINA LEBEDEVA",
    "DMITRY KUZMIN", "ANNA ZAKHAROVA", "YURY LARIN", "TATIANA FROLOVA", "ANDREY LUKIN", "ELENA ZAITSEVA",
    "ALEXEY SMIRNOV", "MARINA PAVLOVA", "ANTON KISELEV", "IRINA KARPOVA", "DMITRIY TITOV", "EKATERINA ZAITSEVA",
    "VLADIMIR LEBEDIN", "OLGA KUZNETSOVA", "NIKOLAI BELIAEV", "ANASTASIA SMIRNOVA", "IGOR ZHDANOV",
    "MARIA BELYAEVA", "ANDREY KURAEV", "EKATERINA ORLOVA", "ALEXANDER PAVLOV", "ANNA KUZMINA", "VLADIMIR RODIN",
    "SVETLANA YAKOVLEVA", "DMITRY DMITRIEV", "YULIA GRIGORIEVA", "VLADIMIR NIKOLAEV", "TATIANA PETROVA",
    "MAXIM KUZNETSOV", "EKATERINA NOVIKOVA", "ANDREI ZUBKOV", "NATALIA MOROZOVA", "SERGEY BORISOV",
    "ALEXANDRA SOKOLOVA", "DMITRY PETROV", "ELENA KONSTANTINOVA", "VLADIMIR GAVRILOV", "ANNA SMIRNOVA",
    "NIKOLAI KISELEV", "MARIA KOMAROVA", "ANDREY TIMOFEEV", "OLGA SMIRNOVA", "VLADIMIR VORONOV",
    "IRINA IVANOVA", "DMITRY EGOROV", "EKATERINA KAZANTSEVA", "ANDREY ZHDANOV", "MARIA SEMENOVA",
    "MAXIM EGOROV", "YULIA KISELEVA", "ALEXANDER KUZNETSOV", "ANASTASIA GAVRILOVA",
    "VLADIMIR KISELEV", "EKATERINA PAVLOVA", "DMITRY BORISOV", "MARIA ZHDANOVA", "ANDREY KUZNETSOV",
    "EKATERINA SERGEEVA", "SERGEY PAVLOV", "TATIANA SHIROKOVA", "ALEXANDER FROLOV", "ANNA PETROVA",
    "NIKOLAI EGOROV", "MARIA KOVALEVSKAYA", "ANDREY ZINOVIEV", "ELENA LUKINA", "DMITRY ZHUKOV",
    "SVETLANA KULIKOVA", "VLADIMIR ZAKHAROV", "IRINA ZHDANOVA", "ANDREI KUZNETSOV", "EKATERINA PONOMAREVA",
    "DMITRY SOKOLOV", "MARIA SHIROKOVA",
    ]
    return random.choice(names)


@app.route('/download_manifest', methods=['GET'])
def download_manifest():
    date = request.args.get('date')
    where_from = request.args.get('where_from')


    # Фильтруем записи в таблице Forms
    filtered_forms = Forms.query.filter(
        Forms.date == date,
        Forms.where_from == where_from,
        Forms.added_to_the_manifest == 'no'
    ).all()

    # Загрузка существующего Excel-файла
    try:
        wb = load_workbook('Sample-Form.xlsx')  # Замените 'Sample-Form.xlsx' на имя вашего существующего файла
        ws = wb.active
    except FileNotFoundError:
        # Обработайте ситуацию, когда файл не найден
        return "Error: Sample-Form.xlsx not found"

    # Обработка данных и запись в Excel
    row_num = ws.max_row + 1  # начинаем с новой строки

    for form in filtered_forms:

        weights = [float(weight) for weight in form.weights.split()]
        price_chance = [15, 20, 25, 10]
        count = 0
        price = random.choice(price_chance)
        vl = 'USD'
        if form.sender_fio:
            if form.sender_fio == 'DAMIR':
                s_n = random_names()
            else:
                s_n = form.sender_fio
        else:
            s_n = random_names()
        purc_count = len(weights)
        
        
        for weight in weights:
            if purc_count != 1:
                count += 1
                if where_from == 'Москва':
                    number = f'{form.city}     {form.number}/{count}'
                else:
                    number = f'{form.city}    0{form.number}/{count}'
            else:
                if where_from == 'Москва':
                    number = f'{form.city}     {form.number}'
                else:
                    number = f'{form.city}    0{form.number}'
            # Добавляем данные в соответствующие столбцы
            ws.cell(row=row_num, column=1, value=s_n.split()[0])  # Имя отправителя
            ws.cell(row=row_num, column=2, value=s_n.split()[-1])  # Фамилия отправителя
            ws.cell(row=row_num, column=3, value='Russian Federation')
            if where_from == 'Москва':
                ws.cell(row=row_num, column=4, value='MOSCOW')
            else:
                ws.cell(row=row_num, column=4, value='S.P.B')
            ws.cell(row=row_num, column=5, value=form.recipient_fio.split()[0])  # Имя получателя
            ws.cell(row=row_num, column=6, value=form.recipient_fio.split()[-1])  # Фамилия получателя
            ws.cell(row=row_num, column=7, value=form.passport)
            ws.cell(row=row_num, column=8, value='Georgia')
            ws.cell(row=row_num, column=9, value=number)
            ws.cell(row=row_num, column=10, value=form.city)
            ws.cell(row=row_num, column=11, value=form.recipient_phone)
            ws.cell(row=row_num, column=12, value=price)
            ws.cell(row=row_num, column=13, value=vl)
            ws.cell(row=row_num, column=14, value=weight)  # Значение веса

            row_num += 1  # Переходим к следующей строке
        form.added_to_the_manifest = 'yes'
        db.session.commit()

    # Сохраняем изменения в файле
    manifest_filename = 'manifest.xlsx'  # Имя нового файла
    wb.save(manifest_filename)

    # Возврат файла для скачивания
    return send_file(manifest_filename, as_attachment=True)




#-------------------------------------------------------------------------------------------------#
# ------------------------------               /expertise           ------------------------------#


@app.route('/expertise')
@login_required
def expertise():

    selected_date = request.args.get('selected_date', None)
    if selected_date:
        latest_date = selected_date
    else:
    # Получаем самую свежую дату из столбца "date"
        latest_date = db.session.query(db.func.max(Expertise.date)).scalar()

    # Получаем все записи с самой свежей датой
    expertise_records = Expertise.query.filter_by(date=latest_date).all()
    unique_dates = db.session.query(func.distinct(Expertise.date)).order_by(desc(Expertise.date)).all()
    unique_dates_list = [date[0].strftime('%Y-%m-%d') for date in unique_dates]
    print(unique_dates_list)
    return render_template('expertise.html', records=expertise_records, date=latest_date, unique_dates_list=unique_dates_list)




@app.route('/expertise_add_record', methods=['POST'])
@login_required
def expertise_add_record():
    try:
        # Получаем данные из POST-запроса
        data = request.get_json()
        # tracking = data['tracking']
        tracking = trecing_redactor(data['tracking'])

        # Загрузка данных из JSON-файла
        with open('expertise_data.json', 'r', encoding='utf-8') as json_file:
            expertise_data = json.load(json_file)

        # Получение записи по ключу data['tracking']
        tracking_key = tracking
        if tracking_key in expertise_data:
            record_data = expertise_data[tracking_key]
            # Дальнейшее использование record_data
        else:
            return jsonify({'error': 'ამანათი არ არსებობს!'}), 404  # Или другой HTTP-статус по вашему выбору

        # Создаем новую запись
        new_record = Expertise(
            Number=data['Number'],
            tracking=tracking,
            comment=data['comment'],
            date=data['date'],
            status=record_data[1],
            weight=record_data[7],
            recipient=record_data[5]
        )

        # Добавляем запись в базу данных
        db.session.add(new_record)
        db.session.commit()

        expertise_list = [{
            'id': new_record.id,
            'status': new_record.status,
            'recipient': new_record.recipient,
            'weight': new_record.weight,
            'Number': new_record.Number,
            'tracking': new_record.tracking,
            'comment': new_record.comment,
            'date': new_record.date.strftime('%Y-%m-%d')
        }]


        
        # expertise_list.append()
        # Ищем дубликаты в JSON-файле
        duplicates = find_duplicates_in_json('expertise_data.json', tracking)
        
        id_trecing = int(new_record.id)

        # Добавляем дубликаты в базу данных
        for duplicate_key in duplicates:
            id_trecing += 1
            dup_data = expertise_data[duplicate_key]
            duplicate_record = Expertise(
                Number=data['Number'],
                tracking=duplicate_key,
                comment='თანაგზავნილი',
                date=data['date'],
                status=dup_data[1],
                weight=dup_data[7],
                recipient=dup_data[5]
            )
            db.session.add(duplicate_record)

            n = {
            'id': id_trecing,
            'status': dup_data[1],
            'recipient': dup_data[5],
            'weight': dup_data[7],
            'Number': data['Number'],
            'tracking': duplicate_key,
            'comment': 'თანაგზავნილი',
            'date': new_record.date.strftime('%Y-%m-%d')
            }
            expertise_list.append(n)

        db.session.commit()
        # Возвращаем данные о новой записи
        return jsonify(expertise_list), 200

    except Exception as e:
        return jsonify({'error': 'Error adding record'}), 500


def find_duplicates_in_json(json_file_path, tracking):
    with open(json_file_path, 'r', encoding='utf-8') as json_file:
        expertise_data = json.load(json_file)

    # Ищем дубликаты
    duplicates = [key for key, value in expertise_data.items()
                  if key != tracking and value[3] == expertise_data[tracking][3] and value[5] == expertise_data[tracking][5]]
    return duplicates
    

@app.route('/expertise_deleted', methods=['POST'])
@login_required
def expertise_deleted():
    try:
        # Получаем данные из POST-запроса
        data = request.get_json()

        # Получаем ID записи, которую нужно удалить
        record_id = data.get('id')

        # Проверяем, что ID был передан
        if record_id is not None:
            # Находим запись в базе данных по ID
            record_to_delete = Expertise.query.get(record_id)

            # Проверяем, что запись существует
            if record_to_delete:
                # Удаляем запись из базы данных
                db.session.delete(record_to_delete)
                db.session.commit()

                return jsonify({'message': 'Record deleted successfully'}), 200
            else:
                return jsonify({'error': 'Record not found'}), 404
        else:
            return jsonify({'error': 'ID not provided'}), 400

    except Exception as e:
        return jsonify({'error': 'Error deleting record'}), 500




from flask import jsonify

@app.route("/rs_xml", methods=["POST"])
@login_required
def rs_xml():
    try:
        if "xmlFile" in request.files:
            xml_file = request.files["xmlFile"]
            date_value = request.form.get('date')

            # Проверка, что файл имеет разрешенное расширение (xml)
            if xml_file and allowed_file(xml_file.filename):
                # Сохранение XML-файла с именем Export.xml в корневую папку
                xml_file.save("Export.xml")

                # Вызов функции xml_convertor после успешной загрузки файла
                xml_convertor()

                # Чтение данных из JSON-файла
                with open('expertise_data.json', 'r', encoding='utf-8') as json_file:
                    expertise_data = json.load(json_file)

                # Создаем список для значений tracking, которые будем обновлять
                tracking_values = list(expertise_data.keys())

                # Обновляем записи в базе данных
                for tracking_key, values in expertise_data.items():
                    update_stmt = (
                        update(Expertise)
                        .where(Expertise.tracking == tracking_key)
                        .values(status=values[1])
                    )
                    db.session.execute(update_stmt)

                # Сохраняем изменения в базе данных только один раз в конце
                db.session.commit()

                # Возвращение JSON-ответа
                return jsonify({'success': True, 'message': 'ბაზა წარმატებით განახლდა'})

        # В случае ошибки
        return jsonify({'success': False, 'message': 'ფაილის ფორმატი არასწორია!'})

    except Exception as e:
        return jsonify({'error': 'Error processing file'}), 500



def status_checker(trecing):
    if len(trecing) < 2:
        return 'არასწორი მონაცემების ფორმატი'

    status_type = trecing[0]
    status_detail = trecing[1]

    if status_type == 'გასატანი' and status_detail == 'დაუბეგრავი':
        return 'დაუბეგრავი!'
    elif status_type == 'გასატანი' and status_detail == 'დაბეგვრადი':
        return 'დაბეგვრადი დასრულებული!'
    elif status_type == 'დასადეკლარირებელი' and status_detail == 'დაბეგვრადი':
        return 'დაბეგვრადი! არ არის მზად!'
    elif status_detail == 'გაურკვეველი':
        return 'ყვითელი!'
    else:
        return 'ამოუცნობიუ სტატუსი'


@app.route('/trecing_checker', methods=['POST'])
def trecing_checker():
    try:
        # Получаем данные из запроса
        trecing_value = trecing_redactor(request.form.get('trecing_value', ''))


        # Читаем содержимое файла expertise_data.json
        with open('expertise_data.json', 'r', encoding='utf-8') as json_file:
            expertise_data = json.load(json_file)

        # Проверяем наличие ключа в объекте
        if trecing_value in expertise_data:
            result = expertise_data[trecing_value]
            information = status_checker(result)
            return jsonify({'success': True, 'result': information})
        else:
            return jsonify({'success': False, 'message': 'Запись не найдена.'})

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})


@app.route('/expertise_export', methods=['POST'])
def expertise_export():
    date = request.form.get('date')

    # Получите данные из базы данных, отфильтрованные по дате
    expertise_data = Expertise.query.filter_by(date=date).all()

    workbook = load_workbook('expertise.xlsx')
    sheet = workbook.active

    start_row = 2  # Starting from the second row as the first row has headers
    col_letters = ['C', 'D', 'E', 'F', 'G', 'H']

    # Устанавливаем стиль шрифта по умолчанию с размером 10
    font = Font(size=10)
    alignment = Alignment(horizontal='center', vertical='center')
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

    for row_idx, data in enumerate(expertise_data, start=start_row):
        for col_idx, col_letter in enumerate(col_letters):
            cell = sheet[f'{col_letter}{row_idx}']
            if col_letter == 'D':
                cell.value = data.recipient
            elif col_letter == 'E':
                cell.value = data.weight
            elif col_letter == 'F':
                cell.value = data.Number
            elif col_letter == 'G':
                cell.value = data.tracking
            elif col_letter == 'H':
                cell.value = data.comment
            # Don't write anything to cell C, just apply formatting
            elif col_letter == 'C':
                pass
            cell.font = font  # Применяем стиль шрифта к ячейке
            cell.alignment = alignment  # Центрируем содержимое ячейки
            cell.border = thin_border  # Добавляем границы

    new_filename = 'expertise-list.xlsx'
    workbook.save(new_filename)
    workbook.close()

    return send_file(new_filename, as_attachment=True)
#-------------------------------------------------------------------------------------------------#
# ------------------------------               end /expertise           ------------------------------#










#-------------------------------------------------------------------------------------------------#
# ------------------------------               other                ------------------------------#
@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))


@app.teardown_request
def teardown_request(exception=None):
    db.session.close()


@app.errorhandler(404)
def page_not_found(error):
    return render_template('error.html', message='Страница не найдена'), 404






# if __name__ == '__main__':
#     socketio.run(app, host='0.0.0.0')

if __name__ == '__main__':
    app.run(host='0.0.0.0', debug=True)


# with app.app_context():
#     db.create_all()




