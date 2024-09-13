from flask import render_template, request, jsonify, send_file, redirect, url_for, flash
from flask.views import MethodView
from flask_login import login_required, current_user
from io import BytesIO
from docx import Document



class DocumentsView(MethodView):
    def __init__(self):
        pass

    decorators = [login_required]

    def get(self):
        access = ['admin', 'Tbilisi']
        if current_user.role not in access:
            flash('თქვენ არ გაქვთ წვდომა ამ გვერდზე', 'error')
            return redirect(url_for('index'))
    
        return render_template('documents.html')

class CreateDocxRequestView(MethodView):
    def __init__(self):
        pass

    decorators = [login_required]

    def post(self):
        # Получаем данные из формы
        ka_name = request.form['ka_f-l_name']
        ge_name = request.form['ge_f-l_name']
        personal_id = request.form['id']
        phone = request.form['phone']
        tracking = request.form['tracking']

        # Открываем существующий документ
        doc = Document('documents/request.docx')

        # Словарь замен
        replacements = {
            "[ka_f/l name]": ka_name,
            "[en f/l_name]": ge_name,
            "[ID]": personal_id,
            "[phone]": phone,
            "[tracking]": f'MP{tracking}'
        }

        # Проходим по каждому параграфу в документе
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    # Заменяем текст
                    paragraph.text = paragraph.text.replace(key, value)

        # Сохраняем документ в память (BytesIO)
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Отправляем документ на скачивание
        file_name = f'{ka_name}.docx'
        return send_file(file_stream, as_attachment=True, download_name=file_name)

class CreateDocxGetTransportingInvoiceView(MethodView):
    def __init__(self):
        pass

    decorators = [login_required]

    def post(self):
        # Получаем данные из формы
        date = request.form['date']
        payer = request.form['payer']
        code = request.form['code']
        tracking = request.form['tracking']
        quantity = request.form['quantity']
        price = request.form['price']
        broker_service = request.form.get('broker_service', 'No')

        tot_transp = int(price) * float(quantity)

        if broker_service == 'on':
            tot_pri = tot_transp + 15

            replacements = {
                "[date]": date,
                "[payer]": payer,
                "[code]": code,
                "[tracking]": f'MP{tracking}',
                "[quantity]": quantity,
                "[price]": price,
                "[tot_transp]": f"{tot_transp:.2f}",
                "[tot_pri]": f"{tot_pri:.2f}",
                "[2]": "2",
                "[საბროკერო მომსახურეობა]": "საბროკერო მომსახურეობა",
                "[ერთ]": "ერთ",
                "[1]": "1",
                "[15]": "15"
            }
        else:
            tot_pri = tot_transp

            replacements = {
                "[date]": date,
                "[payer]": payer,
                "[code]": code,
                "[tracking]": f'MP{tracking}',
                "[quantity]": quantity,
                "[price]": price,
                "[tot_transp]": f"{tot_transp:.2f}",
                "[tot_pri]": f"{tot_pri:.2f}",
                "[2]": "",
                "[საბროკერო მომსახურეობა]": "",
                "[ერთ]": "",
                "[1]": "",
                "[15]": ""
            }
        
        # Открываем существующий документ
        doc = Document('documents/transportation invoice.docx')

        # Функция для замены текста в ячейке
        def replace_text_in_cell(cell, replacements):
            for key, value in replacements.items():
                if key in cell.text:
                    cell.text = cell.text.replace(key, value)

        # Проходим по каждому параграфу в документе
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    # Заменяем текст
                    paragraph.text = paragraph.text.replace(key, value)

        # Проходим по каждой таблице в документе
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_text_in_cell(cell, replacements)

        # Сохраняем документ в память (BytesIO)
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Отправляем документ на скачивание
        file_name = f'Transporting_Invoice_{payer}.docx'
        return send_file(file_stream, as_attachment=True, download_name=file_name)
    


class ManifestCheckerView(MethodView):
    def __init__(self):
        pass

    decorators = [login_required]

    def post(self):
        # Получаем данные из формы
        receiver_first_name_col = request.form.get('receiver_first_name_col', '').strip().split('\n')
        receiver_last_name_col = request.form.get('receiver_last_name_col', '').strip().split('\n')
        weight_col = request.form.get('weight_col', '').strip().split('\n')

        # Убираем символы \r
        receiver_first_name_col = [name.replace('\r', '') for name in receiver_first_name_col]
        receiver_last_name_col = [name.replace('\r', '') for name in receiver_last_name_col]
        weight_col = [weight.replace('\r', '') for weight in weight_col]

        # Проверяем, что все списки одинаковой длины
        if len(receiver_first_name_col) != len(receiver_last_name_col) or len(receiver_first_name_col) != len(weight_col):
            return jsonify({"message": "Data length mismatch", "status": "error"}), 400

        # Создаем словарь из данных и суммируем веса для одинаковых ключей
        data_dict = {}
        for first_name, last_name, weight in zip(receiver_first_name_col, receiver_last_name_col, weight_col):
            key = f"{first_name} {last_name}"
            try:
                weight_value = float(weight)
            except ValueError:
                weight_value = 0  # Если вес не числовой, установим 0

            if key in data_dict:
                data_dict[key] += weight_value
            else:
                data_dict[key] = weight_value

        # Округляем значения до двух знаков после запятой
        data_dict = {key: round(value, 2) for key, value in data_dict.items()}

        # Возвращаем только те записи, где вес >= 29.95
        result = {key: value for key, value in data_dict.items() if value >= 29.95}

        return jsonify({"data": result, "status": "success"})


class PhoneCheckerView(MethodView):
    decorators = [login_required]

    def post(self):
        phone_numbers = request.form.get('phone_numbers', '')  # Получаем текст из textarea
        phone_numbers_list = phone_numbers.splitlines()  # Разделяем по строкам

        invalid_numbers = set()  # Множество для хранения некорректных номеров
        valid_numbers = set()  # Множество для хранения корректных номеров

        # Проверяем каждый номер телефона
        for number in phone_numbers_list:
            number = number.strip()  # Убираем пробелы в начале и конце
            if len(number) == 9 and number.startswith('5') and number.isdigit():
                valid_numbers.add(number)  # Добавляем уникальный номер в множество корректных
            else:
                invalid_numbers.add(number)  # Добавляем уникальный номер в множество некорректных

        # Преобразуем множества обратно в списки для возврата в JSON
        return jsonify(
            valid_numbers=list(valid_numbers),
            invalid_numbers=list(invalid_numbers)
        )





def register_documents_routes(app, db):
    app.add_url_rule('/documents', view_func=DocumentsView.as_view('documents'))
    app.add_url_rule('/create_docxrequest', view_func=CreateDocxRequestView.as_view('create_docxrequest'))
    app.add_url_rule('/create_transporting_invoice', view_func=CreateDocxGetTransportingInvoiceView.as_view('create_transporting_invoice'))
    app.add_url_rule('/manifest_checker', view_func=ManifestCheckerView.as_view('manifest_checker'))
    app.add_url_rule('/phone_number_checker', view_func=PhoneCheckerView.as_view('phone_number_checker'))